from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_el = OxmlElement(f'w:{edge}')
        edge_el.set(qn('w:val'), 'single')
        edge_el.set(qn('w:sz'), '4')
        edge_el.set(qn('w:space'), '0')
        edge_el.set(qn('w:color'), '000000')
        tcBorders.append(edge_el)
    tcPr.append(tcBorders)

def set_run_font(run, font_name='宋体', size=10.5, bold=False, color=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    else:
        run.font.color.rgb = RGBColor(0, 0, 0)

def add_heading_custom(doc, text, level=1):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, '黑体', 16, bold=True)
    else:
        set_run_font(run, '黑体', 14, bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    return p

def add_para(doc, text, align='left', bold=False, size=10.5, first_line_indent=0, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, '宋体', size, bold, color)
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_table_custom(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_border(hdr_cells[i])
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                set_run_font(r, '宋体', 10.5, bold=True)
        if col_widths:
            hdr_cells[i].width = Cm(col_widths[i])
    for row in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
            set_cell_border(row_cells[i])
            for p in row_cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    set_run_font(r, '宋体', 10.5, bold=False)
            if col_widths:
                row_cells[i].width = Cm(col_widths[i])
    return table

doc = Document()

# 页面设置
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)

# 封面
for _ in range(6):
    add_para(doc, "", size=12)
add_para(doc, "2024年第一季度财务分析报告", align='center', bold=True, size=26)
add_para(doc, "", size=12)
add_para(doc, "XX科技股份有限公司", align='center', size=16)
add_para(doc, "报告期间：2024年1月1日 — 2024年3月31日", align='center', size=12)
add_para(doc, "编制日期：2024年4月15日", align='center', size=12)
add_para(doc, "编制部门：财务管理部", align='center', size=12)
add_para(doc, "审计状态：待审计", align='center', size=12, color=(192, 0, 0))
doc.add_page_break()

# 目录
add_heading_custom(doc, "目录", level=1)
for item in [
    "一、财务摘要",
    "二、收入分析", "    2.1 按产品线", "    2.2 按区域",
    "三、成本结构",
    "四、利润分析",
    "五、现金流概况",
    "六、风险提示与建议"
]:
    add_para(doc, item)
doc.add_page_break()

# 一、财务摘要
add_heading_custom(doc, "一、财务摘要", level=1)
add_para(doc, "本季度公司整体财务表现稳健，营业收入同比增长15.9%，净利润同比增长53.7%，毛利率与净利率均实现提升。", first_line_indent=0.74)
add_table_custom(doc,
    ["指标", "2024Q1(万元)", "2023Q1(万元)", "同比变化"],
    [
        ["营业收入", "12,580", "10,850", "+15.9%"],
        ["营业成本", "8,240", "7,320", "+12.6%"],
        ["毛利润", "4,340", "3,530", "+22.9%"],
        ["毛利率", "34.5%", "32.5%", "+2.0pp"],
        ["运营费用", "2,680", "2,450", "+9.4%"],
        ["营业利润", "1,660", "1,080", "+53.7%"],
        ["净利润", "1,245", "810", "+53.7%"],
        ["净利率", "9.9%", "7.5%", "+2.4pp"],
        ["经营现金流", "1,890", "1,420", "+33.1%"],
        ["自由现金流", "980", "650", "+50.8%"],
    ],
    col_widths=[4.5, 3.5, 3.5, 3.0]
)
doc.add_page_break()

# 二、收入分析
add_heading_custom(doc, "二、收入分析", level=1)
add_heading_custom(doc, "2.1 按产品线", level=2)
add_table_custom(doc,
    ["产品线", "2024Q1(万元)", "同比变化", "收入占比"],
    [
        ["智能硬件", "5,340", "+18.1%", "42.4%"],
        ["SaaS订阅", "4,120", "+26.8%", "32.7%"],
        ["咨询服务", "1,860", "+10.7%", "14.8%"],
        ["技术支持", "1,260", "-10.0%", "10.0%"],
    ],
    col_widths=[4.5, 3.5, 3.5, 3.0]
)
add_para(doc, "SaaS订阅业务增长最快（+26.8%），成为第二大收入来源；技术支持收入同比下降10.0%，主要因一次性项目减少。", first_line_indent=0.74)

add_heading_custom(doc, "2.2 按区域", level=2)
add_table_custom(doc,
    ["区域", "2024Q1(万元)", "同比变化", "收入占比"],
    [
        ["华东", "4,650", "+16.8%", "37.0%"],
        ["华南", "3,520", "+18.1%", "28.0%"],
        ["华北", "2,140", "+15.1%", "17.0%"],
        ["西南", "1,380", "+10.4%", "11.0%"],
        ["海外", "890", "+14.1%", "7.1%"],
    ],
    col_widths=[4.5, 3.5, 3.5, 3.0]
)
add_para(doc, "华东与华南合计贡献65%收入，海外业务稳步增长，占比提升至7.1%。", first_line_indent=0.74)
doc.add_page_break()

# 三、成本结构
add_heading_custom(doc, "三、成本结构", level=1)
add_para(doc, "本季度营业成本同比增长12.6%，低于收入增速，规模效应初显。研发投入增长14.0%，持续加大产品创新。", first_line_indent=0.74)
add_table_custom(doc,
    ["成本项目", "2024Q1(万元)", "同比变化", "占总成本比"],
    [
        ["原材料", "3,120", "+9.9%", "37.9%"],
        ["人工成本", "1,860", "+10.7%", "22.6%"],
        ["制造费用", "1,240", "+10.7%", "15.0%"],
        ["研发投入", "980", "+14.0%", "11.9%"],
        ["市场营销", "720", "+10.8%", "8.7%"],
        ["管理费用", "680", "+9.7%", "8.3%"],
        ["折旧摊销", "640", "+16.4%", "7.8%"],
    ],
    col_widths=[4.5, 3.5, 3.5, 3.0]
)
doc.add_page_break()

# 四、利润分析
add_heading_custom(doc, "四、利润分析", level=1)
add_para(doc, "毛利率连续两个季度提升，从2023Q3的31.8%上升至2024Q1的34.5%；净利率同步改善，盈利能力显著增强。", first_line_indent=0.74)
add_table_custom(doc,
    ["季度", "毛利率", "净利率", "营业利润率"],
    [
        ["2023Q1", "32.5%", "7.5%", "10.0%"],
        ["2023Q2", "33.2%", "8.2%", "10.8%"],
        ["2023Q3", "31.8%", "7.1%", "9.5%"],
        ["2023Q4", "34.0%", "9.5%", "12.5%"],
        ["2024Q1", "34.5%", "9.9%", "13.2%"],
    ],
    col_widths=[4.5, 3.5, 3.5, 3.0]
)
doc.add_page_break()

# 五、现金流概况
add_heading_custom(doc, "五、现金流概况", level=1)
add_para(doc, "经营活动净现金流同比增长33.1%，现金流质量良好；投资与筹资活动现金流出处于可控范围，现金净增加额825万元。", first_line_indent=0.74)
add_table_custom(doc,
    ["项目", "2024Q1(万元)", "2023Q1(万元)"],
    [
        ["经营活动现金流入", "4,520", "3,860"],
        ["经营活动现金流出", "2,630", "2,440"],
        ["经营活动净现金流", "1,890", "1,420"],
        ["投资活动净现金流", "-680", "-520"],
        ["筹资活动净现金流", "-420", "-380"],
        ["汇率变动影响", "35", "20"],
        ["现金净增加额", "825", "540"],
    ],
    col_widths=[7.0, 4.0, 4.0]
)
doc.add_page_break()

# 六、风险提示与建议
add_heading_custom(doc, "六、风险提示与建议", level=1)
add_table_custom(doc,
    ["风险类别", "风险等级", "影响描述", "应对建议"],
    [
        ["市场竞争加剧", "中", "SaaS赛道竞争者增多，获客成本上升", "加大产品差异化投入，提升客户留存率"],
        ["原材料价格波动", "中", "芯片及电子元器件价格存在上涨压力", "锁定长期供应协议，优化库存管理"],
        ["汇率波动", "低", "海外收入占比提升，汇率波动影响利润", "开展外汇套期保值，分散结算币种"],
        ["人才流失", "低", "核心技术人才市场竞争激烈", "完善股权激励计划，优化研发文化"],
    ],
    col_widths=[2.5, 1.5, 5.0, 5.0]
)

add_para(doc, "", size=10)
add_para(doc, "【审计声明】", bold=True, size=11)
add_para(doc, "本报告所载财务数据由公司财务管理部编制，基于企业会计准则及内部会计政策。所有重大交易均已入账，报表附注完整。本报告待独立审计师审阅后出具审计意见。", first_line_indent=0.74)

add_para(doc, "", size=10)
add_para(doc, "编制人：___________    审核人：___________    日期：2024年4月15日", align='right', size=10.5)

out_path = "/Users/mashaoguang/Downloads/FireShot/finance_q1_report/2024Q1_财务分析报告.docx"
doc.save(out_path)
print(f"已重新生成兼容版 Word 报告: {out_path}")
